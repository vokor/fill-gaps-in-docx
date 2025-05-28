import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document


def replace_text_in_run(runs, old_text, new_text):
    full_text = "".join([run.text for run in runs])
    if old_text in full_text:
        new_full_text = full_text.replace(old_text, new_text)

        for run in runs:
            run.text = ""

            run.text = new_full_text


def replace_text_in_docx(doc, replacements):
    for p in doc.paragraphs:
        for old_text, new_text in replacements.items():
            for run in p.runs:
                if old_text in run.text:
                    replace_text_in_run(p.runs, old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old_text, new_text in replacements.items():
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if old_text in run.text:
                                replace_text_in_run(p.runs, old_text, new_text)

def upload_file():
    filepath = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if filepath:
        doc_path.set(filepath)

def save_file():
    filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if filepath:
        return filepath
    return None

def process_document():
    if doc_path.get() == "":
        messagebox.showerror("Error", "Please upload a DOCX file.")
        return

    doc = Document(doc_path.get())

    replacements = {
        '{{FIO}}': fio_var.get(),
        '{{INVOICE}}': invoice_var.get(),
        '{{DATE}}': date_var.get(),
        '{{AGREEMENT}}': agreement_var.get()
    }

    replace_text_in_docx(doc, replacements)

    saved_path = save_file()
    if saved_path:
        doc.save(saved_path)
        messagebox.showinfo("Success", f"Document saved as: {saved_path}")

app = tk.Tk()
app.title("DOCX Replacement")

fio_var = tk.StringVar()
invoice_var = tk.StringVar()
date_var = tk.StringVar()
agreement_var = tk.StringVar()
doc_path = tk.StringVar()

tk.Label(app, text="ФИО").grid(row=0, column=0, padx=10, pady=10)
fio_entry = tk.Entry(app, textvariable=fio_var)
fio_entry.grid(row=0, column=1, padx=10, pady=10)

tk.Label(app, text="Счет фактура").grid(row=1, column=0, padx=10, pady=10)
invoice_entry = tk.Entry(app, textvariable=invoice_var)
invoice_entry.grid(row=1, column=1, padx=10, pady=10)

tk.Label(app, text="Дата").grid(row=2, column=0, padx=10, pady=10)
date_entry = tk.Entry(app, textvariable=date_var)
date_entry.grid(row=2, column=1, padx=10, pady=10)

tk.Label(app, text="Договор").grid(row=3, column=0, padx=10, pady=10)
agreement_entry = tk.Entry(app, textvariable=agreement_var)
agreement_entry.grid(row=3, column=1, padx=10, pady=10)

upload_button = tk.Button(app, text="Upload DOCX", command=upload_file)
upload_button.grid(row=4, column=0, columnspan=2, pady=10)

process_button = tk.Button(app, text="Process", command=process_document)
process_button.grid(row=5, column=0, columnspan=2, pady=10)

app.mainloop()

# doc = Document('target.docx')
#
# replacements = {
#     '{{FIO}}': 'Иванов Иван Иванович',
#     '{{INVOICE}}': '№01/147944 от 31.03.2025',
#     '{{DATE}}': '«04» апреля 2025 г.',
#     '{{AGREEMENT}}': '01.06.2012 № 12966.036.1'
# }
#
# replace_text_in_docx(doc, replacements)
#
# doc.save('res.docx')